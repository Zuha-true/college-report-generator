import streamlit as st
from docx import Document
import tempfile, os, subprocess
from helpers.docx_utils import replace_placeholders_in_doc
from helpers.gemini_client import configure, generate_section

# ============ PAGE CONFIG + CUSTOM CSS ============
st.set_page_config(page_title="College Report Generator", layout="centered")
st.markdown(
    """
    <style>
    :root{
        --teal:#0e9aa7;
        --light:#f7fffb;
        --green:#bdecb6;
    }
    .report-header{
        background: linear-gradient(90deg, var(--teal), #16a085);
        color: white;
        padding: 18px;
        border-radius: 8px;
        margin-bottom: 18px;
    }
    .stButton>button {
        background-color: var(--teal);
        color: white;
        border-radius: 8px;
        padding: 8px 16px;
        border: none;
    }
    .stDownloadButton>button {
        background-color: #34c6a5;
        color: white;
        border-radius: 8px;
    }
    </style>
    """, unsafe_allow_html=True
)

st.markdown(
    '<div class="report-header"><h2>📑 College Report Generator</h2>'
    '<div style="font-size:14px">Generate a formatted report from a short project description using Gemini AI</div></div>',
    unsafe_allow_html=True
)

# ============ FORM INPUT ============
with st.form("input_form"):
    st.subheader("🎓 Student & Project Details")

    col1, col2 = st.columns([2, 1])
    with col1:
        college_name = st.text_input("College Name", "ANJUMAN INSTITUTE OF TECHNOLOGY AND MANAGEMENT, BHATKAL")
        affiliation = st.text_input("Affiliation", "Affiliated to Visvesvaraya Technological University, Belagavi")
        department = st.text_input("Department", "Department of Computer Science and Engineering")
        year = st.text_input("Academic Year", "2025 – 2026")
        students = st.text_area("Students (one per line)", 
                                "ARZISH (2AB23CS013)\nNAJMA LANKA (2AB23CS055)\nSHURA SHIPAI (2AB23CS070)\nZUHA RUKNUDDIN (2AB23CS083)")
    with col2:
        project_title = st.text_input("Project Title", "MINI PROJECT (BCS586)")
        professor_name = st.text_input("Professor / Guide", "Prof. Syed Nooreain")
        professor_designation = st.text_input("Designation", "Assistant Professor")

    st.subheader("📝 Project Description")
    project_desc = st.text_area("Short description", 
                                "A database management system to manage student records, including CRUD, search, and reporting features.")

    convert_pdf = st.checkbox("Also produce PDF (requires MS Word on Windows/Mac or LibreOffice on Linux)", value=False)
    st.form_submit_button("Generate report")
    submitted = st.form_submit_button("Generate Report")

# ============ HANDLE SUBMISSION ============
if submitted:
    # --- Get API Key ---
    api_key_input = None
    if "GEMINI_API_KEY" in st.secrets:  # secrets.toml (best for Streamlit Cloud)
        api_key_input = st.secrets["GEMINI_API_KEY"]
    elif os.getenv("GEMINI_API_KEY"):   # environment variable
        api_key_input = os.getenv("GEMINI_API_KEY")
    else:                               # fallback to UI field
        api_key_input = st.text_input("Enter your Gemini API Key to continue", type="password")

    if not api_key_input:
        st.error("❌ No Gemini API Key found. Please set GEMINI_API_KEY env var, add to .streamlit/secrets.toml, or paste above.")
        st.stop()

    # Configure Gemini
    configure(api_key=api_key_input)

    # --- Load template.docx ---
    template_path = "template.docx"
    if not os.path.exists(template_path):
        st.error("❌ template.docx not found. Please create it with placeholders and place it in the project root.")
        st.stop()

    doc = Document(template_path)

    # Replace placeholders
    replace_map = {
        "{COLLEGE_NAME}": college_name,
        "{AFFILIATION}": affiliation,
        "{YEAR}": year,
        "{DEPARTMENT}": department,
        "{STUDENTS}": "\n".join([s.strip() for s in students.splitlines() if s.strip()]),
        "{PROJECT_TITLE}": project_title,
        "{PROFESSOR_NAME}": professor_name,
        "{PROFESSOR_DESIGNATION}": professor_designation,
    }
    doc = replace_placeholders_in_doc(doc, replace_map)

    # --- Generate Report Sections with Gemini ---
    sections = [
        "Abstract",
        "Introduction",
        "Problem Statement",
        "Proposed Solution",
        "Features",
        "Technologies Used",
        "System Requirements",
        "Conclusion",
        "References",
    ]

    st.info("⚡ Generating report content with Gemini...")
    ai_contents = {}
    progress = st.progress(0)
    for i, sec in enumerate(sections):
        ai_contents[sec] = generate_section(project_desc, sec)
        progress.progress((i + 1) / len(sections))

    # Add AI sections to doc
    doc.add_page_break()
    for title, content in ai_contents.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
        doc.add_page_break()

    # Save DOCX
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        docx_path = tmp.name

    with open(docx_path, "rb") as f:
        st.download_button("📥 Download Report (.docx)", f, file_name=f"{project_title.replace(' ','_')}.docx")

    # Optional PDF Conversion
    if convert_pdf:
        try:
            from docx2pdf import convert
            pdf_out = docx_path.replace(".docx", ".pdf")
            convert(docx_path, pdf_out)
            with open(pdf_out, "rb") as f:
                st.download_button("📥 Download Report (.pdf)", f, file_name=f"{project_title.replace(' ','_')}.pdf")
        except Exception as e:
            st.warning("⚠️ docx2pdf failed (MS Word not found). Trying LibreOffice...")
            try:
                pdf_out = docx_path.replace(".docx", ".pdf")
                subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(docx_path), docx_path], check=True)
                if os.path.exists(pdf_out):
                    with open(pdf_out, "rb") as f:
                        st.download_button("📥 Download Report (.pdf)", f, file_name=f"{project_title.replace(' ','_')}.pdf")
            except Exception:
                st.error("❌ PDF conversion failed. Please open the .docx in Word/LibreOffice and export manually.")