# helpers/docx_utils.py
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def replace_placeholders_in_doc(doc: Document, replace_map: dict):
    """
    Replace placeholders across paragraphs and table cells.
    Returns the modified Document object.
    """
    for p in doc.paragraphs:
        for key, val in replace_map.items():
            if key in p.text:
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, val)

    # Also replace inside tables (if your placeholders are in table cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replace_map.items():
                    if key in cell.text:
                        # simplest way: replace text in paragraphs
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, val)
    return doc
